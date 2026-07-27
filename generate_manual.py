#!/usr/bin/env python3
"""Generate Habit Flow user manual (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── Styles ───────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ─── Helper Functions ─────────────────────────────────────────────
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_step(number, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Step {number}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    p.add_run(text)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    p.paragraph_format.left_indent = Cm(1)
    # Add shading
    from docx.oxml.ns import qn
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear'
    })
    shading.append(shd)
    return p

def add_note(text):
    p = doc.add_paragraph()
    r = p.add_run("📌 NOTE: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    p.add_run(text)
    return p

def add_warning(text):
    p = doc.add_paragraph()
    r = p.add_run("⚠️ WARNING: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
    p.add_run(text)
    return p

# ══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()  # spacer
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("🌊 Habit Flow")
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Monthly Habit Tracker — User Manual")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Build better habits, one day at a time.")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run("Version 1.0 — July 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
add_heading_styled("Table of Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. System Requirements",
    "3. Installation & Setup",
    "4. Getting Started (Registration & Login)",
    "5. Dashboard Overview",
    "6. Creating & Managing Habits",
    "7. Tracking Your Daily Progress",
    "8. Understanding the Charts",
    "9. Daily Motivation Quotes",
    "10. Local Data Storage (SQLite)",
    "11. Backups & Data Export",
    "12. Running Locally",
    "13. Deploying Online (Optional)",
    "14. Troubleshooting",
    "15. Frequently Asked Questions",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
add_heading_styled("1. Introduction", level=1)
doc.add_paragraph(
    "Habit Flow is a beautiful, modern habit tracking application designed to help you "
    "build and maintain positive daily habits. With an intuitive glassmorphism interface, "
    "interactive charts, daily motivational quotes, and real-time progress tracking, "
    "Habit Flow makes habit formation enjoyable and insightful."
)
doc.add_paragraph(
    "This application runs entirely on your local machine, storing all your data in a "
    "local SQLite database. You can also deploy it online for free or at minimal cost "
    "to access your habits from anywhere."
)

features = [
    "📊 Three interactive charts — habit distribution, weekly check-ins, monthly progress",
    "🎯 Daily habit tracking with a monthly calendar grid",
    "🔥 Streak tracking with fire badges for motivation",
    "💡 Daily motivational quotes (60 unique quotes)",
    "🎉 Confetti celebration on every check-in",
    "✨ Glassmorphism UI with animated particle background",
    "🔐 Secure login with JWT authentication",
    "📱 Fully responsive — works on desktop, tablet, and mobile",
]
for f in features:
    add_bullet(f)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════
add_heading_styled("2. System Requirements", level=1)
doc.add_paragraph("To run Habit Flow locally, you need:")

reqs = [
    ("Operating System: ", "Windows, macOS, or Linux"),
    ("Node.js: ", "Version 18 or higher"),
    ("npm: ", "Node Package Manager (comes with Node.js)"),
    ("Browser: ", "Any modern browser (Chrome, Firefox, Edge, Safari)"),
    ("Disk Space: ", "~100 MB for the application + dependencies"),
    ("Database: ", "SQLite (bundled with the app, no separate install needed)"),
]
for bold, text in reqs:
    add_bullet(text, bold_prefix=bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  3. INSTALLATION & SETUP
# ══════════════════════════════════════════════════════════════════
add_heading_styled("3. Installation & Setup", level=1)
doc.add_paragraph("Follow these steps to get Habit Flow running on your computer:")

add_step(1, "Install Node.js")
doc.add_paragraph(
    "Download and install Node.js (version 18 or higher) from https://nodejs.org. "
    "This also installs npm (Node Package Manager)."
)

add_step(2, "Download Habit Flow")
doc.add_paragraph("Extract the Habit Flow files to a folder on your computer, e.g.:")
add_code_block("C:\\Users\\YourName\\habit-tracker  (Windows)\n~/Desktop/habit-tracker            (macOS/Linux)")

add_step(3, "Open Terminal / Command Prompt")
doc.add_paragraph("Navigate to the habit-tracker folder:")
add_code_block("cd path/to/habit-tracker")

add_step(4, "Install Dependencies")
doc.add_paragraph("Run the following command to install all required packages:")
add_code_block("npm install")
doc.add_paragraph("This may take 1-2 minutes. You'll see a confirmation when it's done.")

add_step(5, "Start the Server")
doc.add_paragraph("Start the application with:")
add_code_block("PORT=3000 node server.js")
doc.add_paragraph("On Windows, use:")
add_code_block("set PORT=3000 && node server.js")

add_step(6, "Open in Browser")
doc.add_paragraph("Open your web browser and go to:")
add_code_block("http://localhost:3000")
doc.add_paragraph("The Habit Flow login page will appear. You're ready to go!")

add_note("Keep the terminal window open while using the app. Closing it will stop the server.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  4. GETTING STARTED
# ══════════════════════════════════════════════════════════════════
add_heading_styled("4. Getting Started (Registration & Login)", level=1)

add_heading_styled("Creating an Account", level=2)
doc.add_paragraph("When you first open Habit Flow, you'll see the login page.")

add_step(1, 'Click the "Create one" link below the Sign In button')
add_step(2, "Enter a Username (e.g., JaneDoe)")
add_step(3, "Enter your Email address")
add_step(4, "Choose a Password (minimum 6 characters)")
add_step(5, 'Click "Create Account"')

doc.add_paragraph("You'll be automatically logged in and taken to the dashboard.")

add_heading_styled("Logging In", level=2)
doc.add_paragraph("If you already have an account:")
add_step(1, "Enter your Email and Password")
add_step(2, 'Click "Sign In"')

add_note("Your session stays active for 7 days. You won't need to log in again on the same browser.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  5. DASHBOARD OVERVIEW
# ══════════════════════════════════════════════════════════════════
add_heading_styled("5. Dashboard Overview", level=1)
doc.add_paragraph(
    "After logging in, you'll see the main dashboard. Here's a tour of what you'll find:"
)

sections = [
    ("Top Bar", "Shows the Habit Flow logo and your username. Click 'Sign Out' to exit."),
    ("Daily Quote", "A motivational quote that changes every day to keep you inspired."),
    ("Charts Row", "Three interactive charts showing your habit data at a glance."),
    ("Stats Cards", "Four key metrics: Active Habits, Total Check-ins, This Month's count, and Best Streak."),
    ("Your Habits Section", "Lists all your habits with their monthly calendars. Each habit shows its icon, name, description, streak badge, and a color-coded calendar grid."),
    ("+ New Habit Button", "Click to create a new habit."),
]
for title, desc in sections:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}: ")
    r.bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  6. CREATING & MANAGING HABITS
# ══════════════════════════════════════════════════════════════════
add_heading_styled("6. Creating & Managing Habits", level=1)

add_heading_styled("Creating a Habit", level=2)
add_step(1, 'Click the "+ New Habit" button (top right of the habits section)')
add_step(2, "Choose an Icon — Click on an emoji that represents your habit (e.g., 📚 for reading)")
add_step(3, "Choose a Color — Pick a color for your habit's calendar highlights")
add_step(4, "Enter a Habit Name — e.g., 'Morning Meditation'")
add_step(5, "Add a Description (optional) — e.g., '10 minutes every morning'")
add_step(6, 'Click "Create Habit"')

doc.add_paragraph(
    "Your new habit will appear instantly on the dashboard with its own calendar grid."
)

add_heading_styled("Editing a Habit", level=2)
doc.add_paragraph(
    "To edit an existing habit, click the ✏️ (pencil) button on any habit card. "
    "You can change the icon, color, name, or description."
)

add_heading_styled("Deleting a Habit", level=2)
doc.add_paragraph(
    "To delete a habit, click the 🗑️ (trash) button on any habit card. "
    "Confirm the deletion in the popup. Warning: All tracking data for that habit will be permanently lost."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  7. TRACKING YOUR DAILY PROGRESS
# ══════════════════════════════════════════════════════════════════
add_heading_styled("7. Tracking Your Daily Progress", level=1)
doc.add_paragraph(
    "Each habit card contains a monthly calendar grid. Days you've completed are "
    "filled with a colored background."
)

add_heading_styled("Marking a Day as Done", level=2)
add_step(1, "Find the habit you want to track")
add_step(2, "Click on a day number in the calendar grid (past or current days only)")
add_step(3, "A popup will show the habit name and date")
add_step(4, 'Click "Mark Done" to log your check-in')
add_step(5, "🎉 Enjoy the confetti celebration!")

add_heading_styled("Removing a Check-In", level=2)
add_step(1, "Click on a day that is already marked (filled with color)")
add_step(2, 'Click "Remove Check-in" in the popup')
add_step(3, "The day will be unmarked and the stats will update")

add_heading_styled("Understanding the Calendar", level=2)
p = doc.add_paragraph()
r = p.add_run("• ")
p.add_run("Colored days = completed check-ins")
p2 = doc.add_paragraph()
r = p2.add_run("• ")
p2.add_run("Days with a border = today's date")
p3 = doc.add_paragraph()
r = p3.add_run("• ")
p3.add_run("Faded days = future dates (cannot be marked yet)")
p4 = doc.add_paragraph()
r = p4.add_run("• ")
p4.add_run("The counter (e.g., 12/26) shows completed days vs days so far this month")

add_heading_styled("Streak Badges 🔥", level=2)
doc.add_paragraph(
    "When you check in on consecutive days, a fire streak badge appears on your habit card. "
    "It shows your current streak (e.g., '🔥 5-day streak') and glows with a subtle fire animation "
    "to keep you motivated!"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  8. UNDERSTANDING THE CHARTS
# ══════════════════════════════════════════════════════════════════
add_heading_styled("8. Understanding the Charts", level=1)
doc.add_paragraph(
    "Three interactive charts give you visual insights into your habit performance."
)

add_heading_styled("📊 Habit Distribution (Doughnut Chart)", level=2)
doc.add_paragraph(
    "Shows how many check-ins each habit has this month. The size of each colored segment "
    "represents the proportion of total check-ins for that habit. Hover over a segment to "
    "see the exact count."
)

add_heading_styled("📈 Weekly Check-ins (Bar Chart)", level=2)
doc.add_paragraph(
    "Shows the total number of check-ins for each day of the current week (Sunday through Saturday). "
    "This helps you see which days of the week you're most consistent."
)

add_heading_styled("🎯 Monthly Progress (Progress Doughnut)", level=2)
doc.add_paragraph(
    "Shows your overall monthly completion rate. The green portion represents completed check-ins "
    "out of the total possible (habits × days so far). The remaining portion shows what's left to do."
)

add_note("All charts update automatically when you add or remove check-ins.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  9. DAILY MOTIVATION QUOTES
# ══════════════════════════════════════════════════════════════════
add_heading_styled("9. Daily Motivation Quotes", level=1)
doc.add_paragraph(
    "Every day, a new motivational quote appears at the top of your dashboard. "
    "Habit Flow includes 60 hand-picked quotes from famous figures like Mark Twain, "
    "James Clear, Aristotle, Steve Jobs, and many more."
)
doc.add_paragraph(
    "The quote changes automatically each day based on the date, so you'll get a fresh "
    "dose of inspiration every morning. Quotes are displayed with an attribution to "
    "the author."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  10. LOCAL DATA STORAGE
# ══════════════════════════════════════════════════════════════════
add_heading_styled("10. Local Data Storage (SQLite)", level=1)
doc.add_paragraph(
    "Habit Flow stores all your data locally on your computer using SQLite, a powerful "
    "yet lightweight database engine. No internet connection is required once the app is running."
)

add_heading_styled("Where is the Database?", level=2)
doc.add_paragraph(
    "The database file is named data.db and is located in your habit-tracker folder:"
)
add_code_block("habit-tracker/data.db")

add_heading_styled("What Data is Stored?", level=2)
p = doc.add_paragraph()
r = p.add_run("• ")
p.add_run("User accounts (username, email, hashed password)")
p2 = doc.add_paragraph()
r = p2.add_run("• ")
p2.add_run("Habits (name, description, icon, color)")
p3 = doc.add_paragraph()
r = p3.add_run("• ")
p3.add_run("Track records (date, habit, optional note)")

add_heading_styled("Is My Data Secure?", level=2)
doc.add_paragraph(
    "Yes. Passwords are hashed using bcrypt (a strong, industry-standard hashing algorithm) "
    "before being stored. The database file is only accessible on your computer. "
    "If you deploy online, the same security measures apply."
)

add_warning(
    "Do not delete or modify the data.db file directly, as this may corrupt your data. "
    "Always use the app interface to manage your habits."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  11. BACKUPS & DATA EXPORT
# ══════════════════════════════════════════════════════════════════
add_heading_styled("11. Backups & Data Export", level=1)

add_heading_styled("Backing Up Your Data", level=2)
doc.add_paragraph("To back up all your habit data:")
add_step(1, "Make sure the server is stopped (press Ctrl+C in the terminal)")
add_step(2, "Copy the data.db file to a safe location:")
add_code_block("cp habit-tracker/data.db ~/Desktop/habit-backup-2026-07-27.db")
add_step(3, "To restore, copy the backup file back to the habit-tracker folder")

add_heading_styled("Restoring from Backup", level=2)
add_step(1, "Stop the server if it's running")
add_step(2, "Replace data.db with your backup file")
add_step(3, "Restart the server")

add_note("Habit Flow uses SQLite's WAL (Write-Ahead Logging) mode, which ensures data integrity even if the server crashes unexpectedly.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  12. RUNNING LOCALLY
# ══════════════════════════════════════════════════════════════════
add_heading_styled("12. Running Locally", level=1)

add_heading_styled("Standard Start", level=2)
add_code_block("cd habit-tracker")
add_code_block("PORT=3000 node server.js")

add_heading_styled("Running on a Different Port", level=2)
doc.add_paragraph("If port 3000 is already in use, use a different port:")
add_code_block("PORT=8080 node server.js")
add_code_block("PORT=5000 node server.js")
doc.add_paragraph("Then open http://localhost:8080 or http://localhost:5000 in your browser.")

add_heading_styled("Running in Background (Linux/macOS)", level=2)
add_code_block("cd habit-tracker")
add_code_block("nohup node server.js > server.log 2>&1 &")
doc.add_paragraph("This allows the server to keep running even after closing the terminal.")

add_heading_styled("Stopping the Server", level=2)
doc.add_paragraph("Press Ctrl+C in the terminal where the server is running.")
doc.add_paragraph("Or, if running in the background:")
add_code_block("pkill -f 'node server.js'")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  13. DEPLOYING ONLINE (OPTIONAL)
# ══════════════════════════════════════════════════════════════════
add_heading_styled("13. Deploying Online (Optional)", level=1)
doc.add_paragraph(
    "You can deploy Habit Flow to the internet so you can access it from anywhere. "
    "Here are the best free/low-cost options:"
)

add_heading_styled("Option A: Fly.io (Free)", level=2)
doc.add_paragraph("Fly.io offers a generous free tier with 3 shared VMs and 3GB of storage.")
add_step(1, "Create an account at fly.io")
add_step(2, "Install the Fly CLI: curl -L https://fly.io/install.sh | sh")
add_step(3, "Run fly launch in the habit-tracker folder")
add_step(4, "Follow the prompts to deploy")

add_heading_styled("Option B: Railway ($5/month)", level=2)
doc.add_paragraph("Railway offers simple deployment with a $5 starter plan.")
add_step(1, "Push your code to GitHub")
add_step(2, "Create an account at railway.com")
add_step(3, "Click 'New Project' → 'Deploy from GitHub repo'")
add_step(4, "Add environment variable: JWT_SECRET = (your secret key)")

add_heading_styled("Option C: Render ($7/month)", level=2)
add_step(1, "Push your code to GitHub")
add_step(2, "Create an account at render.com")
add_step(3, "Click 'New +' → 'Web Service'")
add_step(4, "Connect your repository")
add_step(5, "Set Start Command: node server.js")

add_warning(
    "When deploying online, ALWAYS set a strong JWT_SECRET environment variable. "
    "Do not use the default secret in production."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  14. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════
add_heading_styled("14. Troubleshooting", level=1)

troubles = [
    ("Port already in use", 
     "Error: listen EADDRINUSE — Another program is using this port.\n"
     "Solution: Use a different port: PORT=3001 node server.js"),
    ("npm install fails",
     "Try: npm install --no-optional or npm cache clean --force then retry.\n"
     "On some systems, you may need to set npm config set strict-ssl false"),
    ("Can't log in",
     "Make sure you registered an account first. Passwords must be at least 6 characters.\n"
     "Try clearing your browser cache and cookies."),
    ("Blank page / nothing loads",
     "Check that the server is running (you should see 'Habit Tracker running' in the terminal).\n"
     "Verify the URL is correct: http://localhost:3000"),
    ("Database errors",
     "If data.db gets corrupted, stop the server, delete data.db, and restart. "
     "You'll need to recreate your account and habits."),
    ("Charts not showing",
     "Charts require Chart.js which loads from CDN. Make sure you have an internet "
     "connection when first loading the page, or the charts won't render."),
]
for title, desc in troubles:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  15. FAQ
# ══════════════════════════════════════════════════════════════════
add_heading_styled("15. Frequently Asked Questions", level=1)

faqs = [
    ("Can I use Habit Flow on my phone?",
     "Yes! The app is fully responsive and works on mobile browsers. "
     "Just open the URL on your phone while the server is running on your computer. "
     "For remote access, deploy it online (see Section 13)."),
    ("Can multiple users use the same installation?",
     "Yes. Each user registers with their own account and sees only their own habits and data. "
     "Users are completely isolated from each other."),
    ("Is my data stored in the cloud?",
     "By default, no. All data is stored locally in the data.db file. "
     "Only if you deploy the app online (Section 13) will the data be on a remote server."),
    ("Can I share my habit data with someone?",
     "You can share your data.db file if you're both running the app locally. "
     "The other person can replace their data.db with yours to see your habits."),
    ("How do I update the app?",
     "Download the latest version and replace the files. Your data.db file will be preserved "
     "if you don't delete it. If there are database schema changes, you may need to start fresh."),
    ("What happens if I forget my password?",
     "Currently, there's no password reset feature. Your admin can delete your user record "
     "from data.db directly (using SQLite Browser), and you can register again."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(f"Q: {q}")
    r.bold = True
    r.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    p2 = doc.add_paragraph(f"A: {a}")
    p2.paragraph_format.space_after = Pt(10)

# ─── Footer ──────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— End of Manual —")
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r.font.italic = True

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer2.add_run("Thank you for using Habit Flow! 🌊\nBuild better habits, one day at a time.")
r.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

# ─── Save ────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Habit_Flow_User_Manual.docx")
doc.save(output_path)
print(f"✅ Manual generated: {output_path}")
